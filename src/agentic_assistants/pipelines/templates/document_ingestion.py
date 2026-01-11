"""
Document ingestion pipeline template.

Pipeline for processing and storing documents (PDF, Office, etc.).
"""

from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Dict, List, Optional

from agentic_assistants.pipelines.pipeline import Pipeline
from agentic_assistants.pipelines.node import Node
from agentic_assistants.utils.logging import get_logger

logger = get_logger(__name__)


@dataclass
class DocumentIngestionConfig:
    """Configuration for document ingestion pipeline."""
    
    input_paths: List[str] = field(default_factory=list)
    collection: str = "documents"
    
    # Supported formats
    supported_formats: List[str] = field(
        default_factory=lambda: [".pdf", ".docx", ".doc", ".txt", ".md", ".rst"]
    )
    
    # Processing
    chunk_size: int = 1024
    chunk_overlap: int = 128
    extract_tables: bool = True
    extract_images: bool = False
    
    # Storage
    store_as_files: bool = True
    store_in_vectordb: bool = True
    artifacts_dir: str = "./data/document_artifacts"


def discover_documents(
    input_paths: List[str],
    supported_formats: List[str],
) -> List[Dict[str, Any]]:
    """
    Discover documents from input paths.
    
    Args:
        input_paths: List of file or directory paths
        supported_formats: Supported file extensions
        
    Returns:
        List of document info dictionaries
    """
    documents = []
    
    for input_path in input_paths:
        path = Path(input_path)
        
        if path.is_file():
            if path.suffix.lower() in supported_formats:
                documents.append({
                    "path": str(path),
                    "name": path.name,
                    "format": path.suffix.lower(),
                    "size_bytes": path.stat().st_size,
                })
        elif path.is_dir():
            for file_path in path.rglob("*"):
                if file_path.is_file() and file_path.suffix.lower() in supported_formats:
                    documents.append({
                        "path": str(file_path),
                        "name": file_path.name,
                        "format": file_path.suffix.lower(),
                        "size_bytes": file_path.stat().st_size,
                    })
    
    logger.info(f"Discovered {len(documents)} documents")
    return documents


def extract_document_content(
    documents: List[Dict[str, Any]],
    extract_tables: bool = True,
    extract_images: bool = False,
) -> List[Dict[str, Any]]:
    """
    Extract content from documents.
    
    Args:
        documents: Output from discover_documents
        extract_tables: Extract tables from documents
        extract_images: Extract images (requires extra processing)
        
    Returns:
        Documents with extracted content
    """
    extracted = []
    
    for doc in documents:
        doc_path = Path(doc["path"])
        doc_format = doc["format"]
        
        try:
            if doc_format == ".pdf":
                content, tables = _extract_pdf(doc_path, extract_tables)
            elif doc_format in [".docx", ".doc"]:
                content, tables = _extract_docx(doc_path, extract_tables)
            elif doc_format in [".txt", ".md", ".rst"]:
                content = doc_path.read_text(encoding="utf-8")
                tables = []
            else:
                content = ""
                tables = []
            
            extracted.append({
                **doc,
                "content": content,
                "tables": tables,
                "success": bool(content),
            })
            
            logger.info(f"Extracted: {doc['name']} ({len(content)} chars)")
            
        except Exception as e:
            logger.error(f"Failed to extract {doc['name']}: {e}")
            extracted.append({
                **doc,
                "content": "",
                "tables": [],
                "success": False,
                "error": str(e),
            })
    
    return extracted


def _extract_pdf(path: Path, extract_tables: bool) -> tuple[str, List[Dict]]:
    """Extract content from PDF."""
    try:
        import pypdf
        
        reader = pypdf.PdfReader(str(path))
        content_parts = []
        
        for page in reader.pages:
            text = page.extract_text()
            if text:
                content_parts.append(text)
        
        content = "\n\n".join(content_parts)
        tables = []  # Table extraction would require additional libraries
        
        return content, tables
        
    except ImportError:
        # Fallback - try pdfplumber
        try:
            import pdfplumber
            
            content_parts = []
            tables = []
            
            with pdfplumber.open(str(path)) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        content_parts.append(text)
                    
                    if extract_tables:
                        page_tables = page.extract_tables()
                        for table in page_tables:
                            if table:
                                tables.append({
                                    "rows": table,
                                    "page": page.page_number,
                                })
            
            return "\n\n".join(content_parts), tables
            
        except ImportError:
            raise ImportError(
                "pypdf or pdfplumber required for PDF extraction. "
                "Install with: pip install pypdf pdfplumber"
            )


def _extract_docx(path: Path, extract_tables: bool) -> tuple[str, List[Dict]]:
    """Extract content from DOCX."""
    try:
        import docx
        
        doc = docx.Document(str(path))
        
        content_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                content_parts.append(para.text)
        
        tables = []
        if extract_tables:
            for table in doc.tables:
                rows = []
                for row in table.rows:
                    cells = [cell.text for cell in row.cells]
                    rows.append(cells)
                if rows:
                    tables.append({"rows": rows})
        
        return "\n\n".join(content_parts), tables
        
    except ImportError:
        raise ImportError(
            "python-docx required for DOCX extraction. "
            "Install with: pip install python-docx"
        )


def process_document_content(
    extracted_documents: List[Dict[str, Any]],
    chunk_size: int = 1024,
    chunk_overlap: int = 128,
) -> List[Dict[str, Any]]:
    """
    Process and chunk document content.
    
    Args:
        extracted_documents: Output from extract_document_content
        chunk_size: Target chunk size
        chunk_overlap: Overlap between chunks
        
    Returns:
        Processed documents with chunks
    """
    from agentic_assistants.data.rag.chunkers import get_chunker, MarkdownChunker
    
    processed = []
    
    for doc in extracted_documents:
        if not doc.get("success") or not doc.get("content"):
            processed.append({**doc, "chunks": []})
            continue
        
        # Choose chunker based on format
        if doc["format"] == ".md":
            chunker = MarkdownChunker(
                chunk_size=chunk_size,
                chunk_overlap=chunk_overlap,
            )
        else:
            chunker = get_chunker(
                "sentence",
                chunk_size=chunk_size,
                chunk_overlap=chunk_overlap,
            )
        
        # Chunk the content
        chunks = chunker.chunk(doc["content"], {
            "source_path": doc["path"],
            "document_name": doc["name"],
            "format": doc["format"],
        })
        
        processed.append({
            **doc,
            "chunks": [c.to_dict() for c in chunks],
            "chunk_count": len(chunks),
        })
        
        logger.info(f"Processed {doc['name']}: {len(chunks)} chunks")
    
    return processed


def augment_document_content(
    processed_documents: List[Dict[str, Any]],
) -> List[Dict[str, Any]]:
    """
    Augment document content with metadata.
    
    Args:
        processed_documents: Output from process_document_content
        
    Returns:
        Augmented documents
    """
    from agentic_assistants.data.rag.augmenters import create_default_augmenter
    from agentic_assistants.data.rag.annotators import create_default_annotator
    from agentic_assistants.data.rag.chunkers import Chunk
    
    augmenter = create_default_augmenter()
    annotator = create_default_annotator()
    
    augmented = []
    
    for doc in processed_documents:
        if not doc.get("chunks"):
            augmented.append(doc)
            continue
        
        context = {
            "source": doc.get("path", ""),
            "source_type": "document",
            "document_name": doc.get("name"),
            "format": doc.get("format"),
        }
        
        # Reconstruct Chunk objects
        chunks = []
        for chunk_dict in doc["chunks"]:
            chunk = Chunk(
                content=chunk_dict["content"],
                index=chunk_dict["index"],
                start_char=chunk_dict["start_char"],
                end_char=chunk_dict["end_char"],
                metadata=chunk_dict.get("metadata", {}),
            )
            chunks.append(chunk)
        
        # Augment and annotate
        chunks = augmenter.augment_batch(chunks, context)
        chunks = annotator.annotate_batch(chunks, context)
        
        doc["chunks"] = [c.to_dict() for c in chunks]
        augmented.append(doc)
    
    return augmented


def store_document_content(
    augmented_documents: List[Dict[str, Any]],
    collection: str,
    store_as_files: bool = True,
    store_in_vectordb: bool = True,
    artifacts_dir: str = "./data/document_artifacts",
) -> Dict[str, Any]:
    """
    Store processed documents.
    
    Args:
        augmented_documents: Output from augment_document_content
        collection: Vector store collection name
        store_as_files: Save as JSON files
        store_in_vectordb: Store in vector database
        artifacts_dir: Directory for file artifacts
        
    Returns:
        Storage statistics
    """
    import json
    import hashlib
    from datetime import datetime
    
    stats = {
        "total_documents": len(augmented_documents),
        "successful_documents": 0,
        "total_chunks": 0,
        "stored_in_vectordb": 0,
        "file_paths": [],
    }
    
    # File storage
    if store_as_files:
        artifacts_path = Path(artifacts_dir) / collection
        artifacts_path.mkdir(parents=True, exist_ok=True)
        
        for doc in augmented_documents:
            if not doc.get("success"):
                continue
            
            stats["successful_documents"] += 1
            stats["total_chunks"] += len(doc.get("chunks", []))
            
            # Create filename
            path_hash = hashlib.sha256(doc["path"].encode()).hexdigest()[:12]
            timestamp = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
            filename = f"{timestamp}_{path_hash}.json"
            filepath = artifacts_path / filename
            
            # Remove large content before saving
            doc_to_save = {k: v for k, v in doc.items() if k != "content"}
            
            with open(filepath, "w") as f:
                json.dump(doc_to_save, f, indent=2)
            
            stats["file_paths"].append(str(filepath))
    
    # Vector store
    if store_in_vectordb:
        try:
            from agentic_assistants.vectordb import LanceDBStore
            from agentic_assistants.config import AgenticConfig
            
            config = AgenticConfig()
            vector_store = LanceDBStore(
                db_path=str(config.vectordb.db_path),
                embedding_model=config.vectordb.embedding_model,
            )
            
            all_documents = []
            all_metadatas = []
            
            for doc in augmented_documents:
                if not doc.get("success"):
                    continue
                
                for chunk in doc.get("chunks", []):
                    all_documents.append(chunk["content"])
                    metadata = chunk.get("metadata", {})
                    metadata["source_path"] = doc["path"]
                    metadata["document_name"] = doc["name"]
                    all_metadatas.append(metadata)
            
            if all_documents:
                vector_store.add(
                    collection=collection,
                    documents=all_documents,
                    metadatas=all_metadatas,
                )
                stats["stored_in_vectordb"] = len(all_documents)
                
        except Exception as e:
            logger.error(f"Failed to store in vector DB: {e}")
    
    return stats


def create_document_ingestion_pipeline(
    config: Optional[DocumentIngestionConfig] = None,
    input_paths: Optional[List[str]] = None,
    collection: str = "documents",
    **kwargs,
) -> Pipeline:
    """
    Create a document ingestion pipeline.
    
    Args:
        config: Pipeline configuration
        input_paths: Document paths (if not using config)
        collection: Vector store collection
        **kwargs: Additional configuration
        
    Returns:
        Configured Pipeline instance
    """
    if config is None:
        config = DocumentIngestionConfig(
            input_paths=input_paths or [],
            collection=collection,
            **kwargs,
        )
    
    # Create nodes
    discover_node = Node(
        func=discover_documents,
        inputs=["input_paths", "supported_formats"],
        outputs=["documents"],
        name="discover_documents",
        tags=["discover", "documents"],
    )
    
    extract_node = Node(
        func=extract_document_content,
        inputs=["documents", "extract_tables", "extract_images"],
        outputs=["extracted_documents"],
        name="extract_document_content",
        tags=["extract", "documents"],
    )
    
    process_node = Node(
        func=process_document_content,
        inputs=["extracted_documents", "chunk_size", "chunk_overlap"],
        outputs=["processed_documents"],
        name="process_document_content",
        tags=["process", "chunk"],
    )
    
    augment_node = Node(
        func=augment_document_content,
        inputs=["processed_documents"],
        outputs=["augmented_documents"],
        name="augment_document_content",
        tags=["augment", "annotate"],
    )
    
    store_node = Node(
        func=store_document_content,
        inputs=["augmented_documents", "collection", "store_as_files", "store_in_vectordb", "artifacts_dir"],
        outputs=["storage_stats"],
        name="store_document_content",
        tags=["store", "vectordb"],
    )
    
    pipeline = Pipeline(
        nodes=[discover_node, extract_node, process_node, augment_node, store_node],
        name="document_ingestion",
        tags=["ingestion", "documents"],
    )
    
    return pipeline
